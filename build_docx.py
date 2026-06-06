# -*- coding: utf-8 -*-
"""Konversi LAPORAN_AKHIR.md -> LAPORAN_AKHIR.docx (python-docx).
Menangani: heading, paragraf, tabel, code block, gambar, list, blockquote,
inline **bold**/`code`/*italic*/[teks](url), rumus $...$ (strip), dan rule ---.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(ROOT, 'LAPORAN_AKHIR.md')

doc = Document()
# Base style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(11)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*|\[[^\]]+\]\([^)]*\))')
IMG_RE = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$')
LINK_RE = re.compile(r'^\[([^\]]+)\]\(([^)]*)\)$')


def shade(cell_or_para_pr, color='F2F2F2'):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), color)
    cell_or_para_pr.append(sh)


def add_inline(p, text, base_bold=False, mono=False):
    text = text.replace('✅', '[v]').replace('✓', '[v]').replace('❌', '[x]')
    text = re.sub(r'\$\$?(.+?)\$\$?', r'\1', text)  # strip math delimiters
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()]); r.bold = base_bold
            if mono: r.font.name = 'Consolas'
        tok = m.group(0)
        if tok.startswith('**'):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        elif tok.startswith('['):
            lm = LINK_RE.match(tok); r = p.add_run(lm.group(1) if lm else tok); r.bold = base_bold
        elif tok.startswith('*'):
            r = p.add_run(tok[1:-1]); r.italic = True; r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); r.bold = base_bold
        if mono: r.font.name = 'Consolas'


def add_code_block(lines):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr())
    cell.paragraphs[0].text = ''
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


_TMP = os.path.join(ROOT, '_imgtmp'); os.makedirs(_TMP, exist_ok=True)

def _downscale(full):
    """Kecilkan resolusi gambar agar ukuran docx ringan; return path siap-embed."""
    try:
        from PIL import Image
        im = Image.open(full)
        if im.mode in ('RGBA', 'P'):
            im = im.convert('RGB')
        cap = 1800
        if max(im.size) > cap:
            ratio = cap / max(im.size)
            im = im.resize((int(im.size[0]*ratio), int(im.size[1]*ratio)), Image.LANCZOS)
        out = os.path.join(_TMP, str(abs(hash(full))) + '.jpg')
        im.save(out, 'JPEG', quality=85, optimize=True)
        return out
    except Exception:
        return full

def add_image(alt, path):
    full = os.path.join(ROOT, path)
    if os.path.exists(full):
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(_downscale(full), width=Inches(4.2))
            return
        except Exception as e:
            pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[GAMBAR: {alt} — {path}]'); r.italic = True; r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)


def make_table(block):
    rows = [r for r in block if r.strip().startswith('|')]
    cells = [[c.strip() for c in r.strip().strip('|').split('|')] for r in rows]
    # drop separator row (---)
    cells = [c for c in cells if not all(set(x) <= set('-: ') and x for x in c)]
    if not cells: return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = t.cell(ri, ci)
            cell.paragraphs[0].text = ''
            add_inline(cell.paragraphs[0], row[ci] if ci < len(row) else '', base_bold=(ri == 0))
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
            if ri == 0:
                shade(cell._tc.get_or_add_tcPr(), 'D9D9D9')
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


with open(MD, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
seen_hr = 0  # cover ends after first ---
while i < len(lines):
    ln = lines[i]
    s = ln.strip()

    # skip raw html div tags
    if s in ('<div align="center">', '</div>'):
        i += 1; continue

    # code fence
    if s.startswith('```'):
        block = []; i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            block.append(lines[i]); i += 1
        i += 1
        add_code_block(block); continue

    # horizontal rule
    if s == '---':
        seen_hr += 1
        doc.add_paragraph().add_run('_' * 60).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        i += 1; continue

    # image
    im = IMG_RE.match(s)
    if im:
        add_image(im.group(1), im.group(2)); i += 1; continue

    # table block
    if s.startswith('|'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i]); i += 1
        make_table(block); continue

    # headings
    if s.startswith('### '):
        h = doc.add_heading(level=3); add_inline(h, s[4:]); i += 1; continue
    if s.startswith('## '):
        h = doc.add_heading(level=2)
        if seen_hr == 0: h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(h, s[3:]); i += 1; continue
    if s.startswith('# '):
        h = doc.add_heading(level=1)
        if seen_hr == 0: h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(h, s[2:]); i += 1; continue

    # blockquote
    if s.startswith('>'):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        r0 = p.add_run('NOTE: '); r0.bold = True; r0.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        add_inline(p, s.lstrip('> ').strip()); i += 1; continue

    # bullet list (- or *) with nesting
    if re.match(r'^(\s*)[-*] ', ln):
        indent = len(ln) - len(ln.lstrip())
        style = 'List Bullet' if indent < 2 else 'List Bullet 2'
        p = doc.add_paragraph(style=style)
        add_inline(p, s[2:]); i += 1; continue

    # numbered list
    if re.match(r'^\s*\d+\. ', ln):
        p = doc.add_paragraph(style='List Number')
        add_inline(p, re.sub(r'^\s*\d+\.\s*', '', ln)); i += 1; continue

    # blank
    if s == '':
        i += 1; continue

    # normal paragraph (center cover lines)
    p = doc.add_paragraph()
    if seen_hr == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, s)
    i += 1

out = os.path.join(ROOT, 'LAPORAN_AKHIR.docx')
doc.save(out)
print('OK ->', out)
